"""
Markdown -> PDF and DOCX.

The Markdown produced by packaging/markdown.py is the source of truth; both
renderers consume it, so the three formats cannot drift apart.

Both backends are optional imports. A missing library disables that one
format with a clear message rather than breaking package generation --
Markdown always works.

    PDF   reportlab      (pure Python, no system dependencies)
    DOCX  python-docx
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_TABLE_ROW = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\|[\s:\-|]+\|\s*$")
_RULE = re.compile(r"^---+\s*$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_HTML_COMMENT = re.compile(r"^<!--.*-->\s*$")

_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
_CODE = re.compile(r"`([^`]+)`")


class RendererUnavailable(RuntimeError):
    """The optional library for a requested format is not installed."""


@dataclass
class Block:
    kind: str          # heading | para | quote | table | rule | bullet
    text: str = ""
    level: int = 0
    rows: list = None


# ---------------------------------------------------------------------------
# Markdown -> blocks
# ---------------------------------------------------------------------------


def parse_blocks(md: str) -> list[Block]:
    """Parse the subset of Markdown that render_markdown() emits."""
    blocks: list[Block] = []
    lines = md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or _HTML_COMMENT.match(stripped):
            i += 1
            continue

        if _RULE.match(stripped):
            blocks.append(Block(kind="rule"))
            i += 1
            continue

        heading = _HEADING.match(stripped)
        if heading:
            blocks.append(Block(kind="heading", level=len(heading.group(1)), text=heading.group(2).strip()))
            i += 1
            continue

        if _TABLE_ROW.match(stripped):
            rows, i = _consume_table(lines, i)
            if rows:
                blocks.append(Block(kind="table", rows=rows))
            continue

        if stripped.startswith(">"):
            quote, i = _consume_quote(lines, i)
            blocks.append(Block(kind="quote", text=quote))
            continue

        bullet = _BULLET.match(stripped)
        if bullet:
            blocks.append(Block(kind="bullet", text=bullet.group(1).strip()))
            i += 1
            continue

        para, i = _consume_paragraph(lines, i)
        if para.strip():
            blocks.append(Block(kind="para", text=para.strip()))

    return blocks


def _consume_table(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while i < len(lines) and _TABLE_ROW.match(lines[i].strip()):
        raw = lines[i].strip()
        if not _TABLE_SEP.match(raw):
            cells = [c.strip() for c in raw.strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def _consume_quote(lines: list[str], i: int) -> tuple[str, int]:
    parts: list[str] = []
    while i < len(lines) and lines[i].strip().startswith(">"):
        parts.append(lines[i].strip().lstrip(">").strip())
        i += 1
    return "\n".join(parts).strip(), i


def _consume_paragraph(lines: list[str], i: int) -> tuple[str, int]:
    parts: list[str] = []
    while i < len(lines):
        stripped = lines[i].strip()
        if (
            not stripped
            or _HEADING.match(stripped)
            or _TABLE_ROW.match(stripped)
            or _RULE.match(stripped)
            or stripped.startswith(">")
            or _BULLET.match(stripped)
            or _HTML_COMMENT.match(stripped)
        ):
            break
        parts.append(stripped)
        i += 1
    return " ".join(parts), i


def strip_inline(text: str) -> str:
    """Plain text with Markdown emphasis and link syntax removed."""
    text = _LINK.sub(r"\1", text)
    text = _BOLD.sub(r"\1", text)
    text = _ITALIC.sub(r"\1", text)
    text = _CODE.sub(r"\1", text)
    return text


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def render_pdf(md: str, output_path: Path) -> Path:
    """Render Markdown to PDF via reportlab."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            HRFlowable,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RendererUnavailable(
            "PDF export needs reportlab. Install it with: pip install reportlab"
        ) from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PkgBody", parent=styles["BodyText"], fontSize=9.5, leading=13.5, alignment=TA_LEFT
    )
    quote = ParagraphStyle(
        "PkgQuote",
        parent=body,
        leftIndent=10,
        borderPadding=4,
        textColor=colors.HexColor("#333333"),
        backColor=colors.HexColor("#F4F4F4"),
        fontName="Helvetica",
        fontSize=9,
        leading=12.5,
    )
    heading_styles = {
        1: ParagraphStyle("PkgH1", parent=styles["Heading1"], fontSize=18, spaceAfter=8),
        2: ParagraphStyle("PkgH2", parent=styles["Heading2"], fontSize=14, spaceBefore=12, spaceAfter=6),
        3: ParagraphStyle("PkgH3", parent=styles["Heading3"], fontSize=12, spaceBefore=10, spaceAfter=4),
        4: ParagraphStyle("PkgH4", parent=styles["Heading4"], fontSize=10.5, spaceBefore=8, spaceAfter=3),
        5: ParagraphStyle("PkgH5", parent=styles["Heading5"], fontSize=10, spaceBefore=6, spaceAfter=3),
        6: ParagraphStyle("PkgH6", parent=styles["Heading6"], fontSize=9.5, spaceBefore=6, spaceAfter=3),
    }

    story = []
    for block in parse_blocks(md):
        if block.kind == "heading":
            style = heading_styles.get(block.level, heading_styles[6])
            story.append(Paragraph(_to_rl(block.text), style))
        elif block.kind == "para":
            story.append(Paragraph(_to_rl(block.text), body))
            story.append(Spacer(1, 3))
        elif block.kind == "bullet":
            story.append(Paragraph("• " + _to_rl(block.text), body))
        elif block.kind == "quote":
            for chunk in block.text.split("\n\n"):
                if chunk.strip():
                    story.append(Paragraph(_to_rl(chunk.replace("\n", "<br/>")), quote))
                    story.append(Spacer(1, 3))
        elif block.kind == "rule":
            story.append(Spacer(1, 5))
            story.append(HRFlowable(width="100%", color=colors.HexColor("#CCCCCC")))
            story.append(Spacer(1, 5))
        elif block.kind == "table" and block.rows:
            story.append(_rl_table(block.rows, body, Table, TableStyle, colors, mm))
            story.append(Spacer(1, 6))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Knowledge Transfer Package",
    )
    doc.build(story)
    return output_path


def _rl_table(rows, body_style, Table, TableStyle, colors, mm):
    from reportlab.platypus import Paragraph

    cell_style = body_style.clone("PkgCell")
    cell_style.fontSize = 8.5
    cell_style.leading = 11

    data = [[Paragraph(_to_rl(cell), cell_style) for cell in row] for row in rows]
    table = Table(data, hAlign="LEFT", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _to_rl(text: str) -> str:
    """Markdown inline -> reportlab's mini-HTML, escaping the rest."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = _LINK.sub(r'<link href="\2" color="blue">\1</link>', text)
    text = _BOLD.sub(r"<b>\1</b>", text)
    text = _ITALIC.sub(r"<i>\1</i>", text)
    text = _CODE.sub(r'<font face="Courier">\1</font>', text)
    return text


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def render_docx(md: str, output_path: Path) -> Path:
    """Render Markdown to DOCX via python-docx."""
    try:
        import docx
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RendererUnavailable(
            "DOCX export needs python-docx. Install it with: pip install python-docx"
        ) from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = docx.Document()

    for block in parse_blocks(md):
        if block.kind == "heading":
            document.add_heading(strip_inline(block.text), level=min(block.level, 4))

        elif block.kind == "para":
            paragraph = document.add_paragraph()
            _docx_runs(paragraph, block.text, Pt, RGBColor)

        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            _docx_runs(paragraph, block.text, Pt, RGBColor)

        elif block.kind == "quote":
            for chunk in block.text.split("\n"):
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Pt(18)
                run = paragraph.add_run(strip_inline(chunk))
                run.italic = True
                run.font.size = Pt(9)

        elif block.kind == "rule":
            document.add_paragraph("_" * 60)

        elif block.kind == "table" and block.rows:
            _docx_table(document, block.rows)

    document.save(str(output_path))
    return output_path


def _docx_table(document, rows) -> None:
    columns = max(len(r) for r in rows)
    table = document.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(columns):
            value = strip_inline(row[column]) if column < len(row) else ""
            cells[column].text = value
            if index == 0:
                for paragraph in cells[column].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _docx_runs(paragraph, text: str, Pt, RGBColor) -> None:
    """Add text to a paragraph, honouring bold/italic/link markup."""
    token = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))")
    for part in token.split(text):
        if not part:
            continue
        link = _LINK.fullmatch(part)
        if link:
            run = paragraph.add_run(link.group(1))
            run.font.color.rgb = RGBColor(0x1A, 0x0D, 0xAB)
            run.underline = True
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            continue
        if part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
            continue
        paragraph.add_run(part)


# ---------------------------------------------------------------------------


def available_formats() -> dict[str, bool]:
    """Which optional renderers are installed right now."""
    formats = {"md": True, "pdf": False, "docx": False}
    try:
        import reportlab  # noqa: F401

        formats["pdf"] = True
    except ImportError:
        pass
    try:
        import docx  # noqa: F401

        formats["docx"] = True
    except ImportError:
        pass
    return formats
